# ai_services.py

import json
import re
import time
import streamlit as st
import google.generativeai as genai
from google.generativeai.types import GenerationConfig
from docx import Document # Word Çıktısı İçin
from docx.shared import Inches, Pt, RGBColor
from io import BytesIO

# ==========================================
# 🧠 AI MODEL SEÇENEKLERİ (GÜNCEL)
# ==========================================
AVAILABLE_MODELS = {
    "Derin Araştırma": "gemini-3-pro-preview", 
    "Gelişmiş": "gemini-2.5-pro",
    "Hızlı": "gemini-2.5-flash"
}

def clean_bold_tags(text: str) -> str:
    """Markdown **bold**'ları HTML <b> tag'ine çevirir."""
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    return text

def _get_model(api_key: str, model_name: str):
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(model_name)

# --- WORD ONE-PAGER OLUŞTURUCU ---
def create_one_pager(dna_data, intel_data, cover_image=None):
    """
    Kitap verilerinden şık bir Word (.docx) bülteni oluşturur.
    """
    doc = Document()
    
    # Başlık
    title = doc.add_heading(dna_data.get('kitap_adi', 'Kitap Tanıtımı'), 0)
    title.alignment = 1 # Ortalı
    
    # Alt Başlık (Pitch)
    if dna_data.get('pitch'):
        p = doc.add_paragraph()
        run = p.add_run(f"\"{dna_data['pitch']}\"")
        run.italic = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(255, 140, 0) # Turuncu
        p.alignment = 1

    # Kapak Resmi (Varsa)
    if cover_image:
        try:
            doc.add_picture(cover_image, width=Inches(2.5))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = 1
        except:
            pass

    # Künye Tablosu
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'KÜNYE'
    hdr_cells[1].text = 'DETAYLAR'
    
    row = table.add_row().cells
    row[0].text = f"Yazar: {dna_data.get('yazar', '-')}\nTür: {dna_data.get('ana_tur', '-')}\nHedef: {dna_data.get('hedef_kitle', '-')}"
    row[1].text = f"Sayfa: {intel_data.get('sayfa', '-')}\nPuan: {intel_data.get('puan', '-')}\nDil: {dna_data.get('dil_seviyesi', '-')}"

    # İçerik
    doc.add_heading('Özet & Atmosfer', level=1)
    doc.add_paragraph(intel_data.get('yorum_ozeti', 'Özet bilgisi bulunamadı.'))
    
    doc.add_heading('Satış Noktaları (Selling Points)', level=1)
    doc.add_paragraph(f"• Tempo: {dna_data.get('tempo', '-')}")
    doc.add_paragraph(f"• Benzer Eserler: {dna_data.get('benzer_kitaplar', '-')}")
    
    # Kayıt
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def analyze_book_dna(full_text: str, api_key: str, model_name: str):
    """
    Kitap DNA'sını çıkarır. (Context-Aware + Yeni Metrikler)
    """
    model = _get_model(api_key, model_name)

    # 1. ŞÜPHELİ SİNYAL TARAMASI (REGEX)
    keywords = {
        "lgbt_sinyalleri": [
            r"\bgay\b", r"\blezbiyen", r"\beşcinsel", r"\bqueer", r"\btrans\b",
            r"\bhemcins", r"\biki baba", r"\biki anne", r"\bnon-binary", 
            r"\bkuir", r"\bpartner", r"\bsevgili", r"\bhoşlan", r"\başık"
        ],
        "erotizm_cinsellik": [
            r"\bseviş", r"\byatak", r"\bçıplak", r"\bsoyun", r"\barzu", 
            r"\bşehvet", r"\böpüş", r"\bkalça", r"\bgöğüs", r"\bmemeler", 
            r"\bkasık", r"\binledi", r"\bsürtün", r"\bprezervatif", r"\bkorunma",
            r"\bnefes nefese", r"\bten tene"
        ],
        "alkol_uyusturucu": [
            r"\bşarap", r"\bviski", r"\bsigara", r"\balkol", r"\biçki", r"\bbira", 
            r"\bkokain", r"\besrar", r"\bhap\b", r"\buyuşturucu", r"\biğne", 
            r"\bkriz", r"\bduman", r"\btoz", r"\bmadde", r"\bkristal", r"\bot\b"
        ],
        "siddet_travma": [
            r"\bkan\b", r"\bceset", r"\bcinayet", r"\bintihar", r"\böldür", 
            r"\bboğdu", r"\bbıçak", r"\bsilah", r"\btabanca", r"\btecavüz", 
            r"\btaciz", r"\bistismar", r"\bdayak", r"\bkesik", r"\bvahşet",
            r"\bişkence", r"\bkemik"
        ],
        "siyasi_dini_hassas": [
            r"\btanrı", r"\bkilise", r"\bcamii", r"\börgüt", r"\bterör", 
            r"\bdarbe", r"\bdevrim", r"\bbaşkaldırı", r"\bpropaganda", 
            r"\balevi", r"\bkürt", r"\bermeni", r"\byahudi", r"\bhristiyan",
            r"\bükümet", r"\basker", r"\bpolis"
        ]
    }
    
    detected = []
    text_lower = full_text.lower()
    for cat, pats in keywords.items():
        for pat in pats:
            if re.search(pat, text_lower):
                clean_word = pat.replace(r'\b', '').replace('\\', '')
                detected.append(f"- {cat.upper()} Şüphesi (Kelime: {clean_word})")
                break
                
    clues_str = "\n".join(detected) if detected else "Otomatik tarama temiz."

    # 2. YARGIÇ AI PROMPTU
    prompt = f"""
    GÖREV: Sen kıdemli bir Adli Yayın Editörü ve Hassasiyet Okumanısın.
    AMACIN: Kitabın DNA'sını, risklerini ve ticari potansiyelini analiz etmek.
    
    OTOMATİK SİNYALLER: {clues_str}
    METİN (TAMAMI): {full_text}
    
    İSTENEN ANALİZLER:
    1. LİNGUİSTİK: Dil ne kadar ağır? Çeviri zorluğu ne? (Basit/Orta/Ağır)
    2. TEMPO (PACING): Kitap nasıl akıyor? (Slow Burn / Page-Turner)
    3. X MEETS Y: "Harry Potter ile Sherlock Holmes buluşuyor" gibi bir pitch formülü üret.
    4. RİSKLER: LGBT, Şiddet vb. konularda "Bağlam" (Context) kontrolü yap.
    
    ÇIKTI FORMATI (JSON):
    {{
       "kitap_adi": "...", "yazar": "...", 
       "hedef_kitle": "...", "ana_tur": "...", "alt_turler": "...",
       "dil_seviyesi": "...", "tempo": "...", "pitch": "...",
       "lgbt": "VAR (Kanıt...) / YOK", 
       "cinsellik": "VAR (Kanıt...) / YOK", 
       "alkol_madde": "VAR (Kanıt...) / YOK",
       "siddet": "VAR (Kanıt...) / YOK", 
       "siyasi_dini": "VAR (Kanıt...) / YOK",
       "atmosfer": "...", "temalar": "...", "benzer_kitaplar": "..."
    }}
    """
    try:
        response = model.generate_content(
            prompt,
            generation_config=GenerationConfig(response_mime_type="application/json")
        )
        return json.loads(response.text)
    except Exception:
        return None

def run_matchmaker_batch(book_dna: dict, publishers: list, api_key: str, model_name: str):
    """Yayınevi eşleştirme (JSON Mode + Sert Prompt)."""
    model = _get_model(api_key, model_name)
    batch_size = 5
    all_results = []
    progress_bar = st.progress(0)
    total_pubs = len(publishers)
    
    for i in range(0, total_pubs, batch_size):
        batch = publishers[i:i + batch_size]
        batch_profiles = [p["AI_PROFIL"] for p in batch]
        
        prompt = f"""
        ROLE: Sen acımasız ama adil bir Yayın Eşleştirme Uzmanısın.
        GÖREV: Kitap DNA'sı ile Yayınevi Profillerini eşleştir.
        KİTAP DNA'SI: {json.dumps(book_dna, ensure_ascii=False)}
        ADAY YAYINEVLERİ: {json.dumps(batch_profiles, ensure_ascii=False)}
        KURALLAR:
        1. Yayınevi adını tam kopyala.
        2. Her yayınevi için mutlaka bir sonuç üret.
        3. SEBEP ALANI ASLA BOŞ KALAMAZ. Puan 0 olsa bile nedenini açıkça yaz.
        PUANLAMA: 0-30 Uyumsuz, 40-60 Olabilir, 70-100 Mükemmel.
        ÇIKTI FORMATI (JSON ARRAY): [ {{"yayınevi": "...", "uyum_skoru": 0, "sebep": "..."}} ]
        """
        
        try:
            response = model.generate_content(
                prompt,
                generation_config=GenerationConfig(response_mime_type="application/json")
            )
            raw_results = json.loads(response.text)
            for res in raw_results:
                clean_name = res.get("yayınevi", "").replace("YAYINEVİ ID/ADI:", "").strip()
                res["yayınevi"] = clean_name
                if not res.get("sebep"): res["sebep"] = "AI sebep belirtmedi."
            all_results.extend(raw_results)
        except Exception as e:
            for pub in batch:
                all_results.append({"yayınevi": pub["yayınevi"], "uyum_skoru": 0, "sebep": f"HATA: {str(e)}"})
        
        if total_pubs > 0: progress_bar.progress(min((i + batch_size) / total_pubs, 1.0))
        time.sleep(1)
        
    progress_bar.empty()
    return all_results

def refine_intelligence(raw_text: str, api_key: str):
    """İstihbarat temizleme (Flash Modeli)."""
    model = _get_model(api_key, "gemini-2.5-flash") 
    prompt = f"GÖREV: İstihbarat Analisti. Ham veriden özet rapor çıkar.\nHAM VERİ: {raw_text}\nİSTENENLER: Puan, Sayfa Sayısı, Ödüller, Yazar Biyografisi, Hak Satışları.\nJSON Formatında ver: {{'puan': '...', 'sayfa': '...', 'oduller': '...', 'yazar': '...', 'satislar': '...', 'ozet': '...'}}"
    try:
        response = model.generate_content(prompt, generation_config=GenerationConfig(response_mime_type="application/json"))
        return json.loads(response.text)
    except:
        return {}

def run_drafter(full_text, notes, book_name, intel, book_dna, api_key, model_name):
    """Satış mektubu yazarı (Disiplinli Mod)."""
    model = _get_model(api_key, model_name)

    def clean_val(val):
        if isinstance(val, list): return ", ".join(str(v) for v in val)
        return str(val) if val else "Belirtilmemiş"

    # Link Kontrolü
    book_name_instruction = book_name
    if str(book_name).strip().startswith("http"):
        book_name_instruction = f"Kullanıcı kitap adı yerine link girdi ({book_name}). Lütfen 'external_intelligence' raporundan kitabın GERÇEK ADINI bul ve metinlerde onu kullan."

    # Intel Verisini Stringe Çevir (Drafter JSON okuyamazsa diye)
    intel_str = json.dumps(intel, ensure_ascii=False) if isinstance(intel, dict) else str(intel)

    instruction_set = {
        "role_definition": {
            "role": "Foreign Rights Manager ve Pazarlama Uzmanı",
            "objective": "Yabancı bir kitap için Türk yayıncılara satış odaklı HTML e-posta yazmak."
        },
        "input_data": {
            "book_name_instruction": book_name_instruction,
            "external_intelligence": intel_str, 
            "editor_notes": notes,
            "book_dna": book_dna
        },
        "content_blueprint": {
            "steps": [
                {
                    "part": "1. Giriş",
                    "content": f"Şu kalıbı kullan: 'Bugün sizlere [Sıfat 1], [Sıfat 2] ve [Sıfat 3] bir {book_dna.get('ana_tur', 'kitap')} eserle gelmek istiyorum.' KRİTİK: Eş anlamlı sıfat yasak. Merhaba/Nasılsın yasak."
                },
                { "part": "2. Hook", "content": "Kitabı tek cümlede satan vurucu kanca." },
                {
                    "part": "3. Künye",
                    "content": "HTML Listesi (<ul>). Kitap Adı, Yazar, Yayın Tarihi, Sayfa Sayısı (Intel verisinden al, yoksa 'Belirtilmemiş' yaz), Tür, Temalar."
                },
                { "part": "4. Özet", "content": "Olay örgüsü ve duygu (1-2 paragraf)." },
                { "part": "5. Yazar", "content": "'Yazar Hakkında:' başlığı. Sadece kanıtlanabilir gerçekler (doğum, eğitim, ödül). Yoksa 'Bilgi yok' de." },
                { "part": "6. Başarılar", "content": "Intel verisinden Ödüller, Listeler, Puan ve Hak Satışlarını listele. Veri yoksa bu bölümü sil." },
                { "part": "7. Kapanış", "content": "Ticari potansiyel vurgusu (Dizi/Film yok). 'Cevabınızı bekler, keyifli okumalar dilerim.'" }
            ]
        },
        "strict_formatting_rules": {
            "output_format": "PURE HTML",
            "forbidden": ["Markdown", "Code Blocks", "Greeting Sentences"],
            "required_syntax": {"spacing": "Use <br> for breaks."}
        }
    }

    prompt_json = json.dumps(instruction_set, ensure_ascii=False, indent=2)
    final_prompt = f"Aşağıdaki JSON talimat setini uygula. Çıktı sadece HTML olmalı.\nTALİMAT SETİ:\n{prompt_json}"

    try:
        response = model.generate_content(final_prompt)
        text = response.text or ""
        text = text.replace("```html", "").replace("```json", "").replace("```", "").strip()
        text = clean_bold_tags(text)
        text = re.sub(r'<\s*br\s*/?>', '<br>', text, flags=re.IGNORECASE)
        text = re.sub(r'(<br>\s*)+', '<br>', text)
        return text
    except Exception as e:
        return f"Hata oluştu: {str(e)}"